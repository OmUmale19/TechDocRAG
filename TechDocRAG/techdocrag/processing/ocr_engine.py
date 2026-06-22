"""
Multi-engine OCR system with fallback capabilities.
Supports Tesseract, Azure Cognitive Services, and AWS Textract.
"""

import asyncio
import io
from pathlib import Path
from typing import Dict, Any, Tuple, Optional, List
from PIL import Image
import pytesseract
import fitz  # PyMuPDF

from ..core.config import OCRConfig
from ..utils.logging import get_logger, performance_log
from ..utils.exceptions import OCRError
from ..utils.helpers import retry_on_failure

logger = get_logger(__name__)


class OCREngine:
    """
    Intelligent OCR engine with multiple backends and fallback capabilities.
    
    Architecture Decision: Multi-engine approach for maximum reliability
    - Tesseract: Free, reliable baseline
    - Azure: Enterprise-grade accuracy
    - AWS Textract: Advanced layout understanding
    - EasyOCR: Deep learning fallback
    """
    
    def __init__(self, config: OCRConfig):
        self.config = config
        self.engines = {
            'tesseract': self._tesseract_ocr,
            'azure': self._azure_ocr,
            'aws': self._aws_ocr,
            'easyocr': self._easyocr_ocr
        }
        
        # Initialize available engines
        self.available_engines = self._check_available_engines()
        logger.info(f"OCR engines available: {list(self.available_engines.keys())}")
    
    @performance_log("ocr_extraction")
    async def extract_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from document using best available OCR engine.
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (extracted_text, metadata)
            
        Raises:
            OCRError: If all OCR engines fail
        """
        logger.info(f"Starting OCR extraction: {file_path}")
        
        # Determine document type and prepare for OCR
        doc_type = self._get_document_type(file_path)
        
        if doc_type == 'pdf':
            pages_data = await self._extract_pdf_pages(file_path)
        elif doc_type == 'image':
            pages_data = [{'image': Image.open(file_path), 'page_num': 1}]
        elif doc_type == 'docx':
            # For DOCX, extract text directly without OCR
            return await self._extract_docx_text(file_path)
        else:
            raise OCRError(f"Unsupported document type: {doc_type}")
        
        # Try OCR engines in order of preference
        engines_to_try = [self.config.primary_engine] + self.config.fallback_engines
        
        for engine_name in engines_to_try:
            if engine_name not in self.available_engines:
                logger.warning(f"OCR engine not available: {engine_name}")
                continue
                
            try:
                logger.info(f"Attempting OCR with {engine_name}")
                text, metadata = await self._run_ocr_engine(engine_name, pages_data)
                
                # Validate extraction quality
                if self._validate_extraction(text, metadata):
                    metadata.update({
                        'ocr_engine': engine_name,
                        'total_pages': len(pages_data),
                        'extraction_method': 'ocr'
                    })
                    logger.info(f"OCR successful with {engine_name}: {len(text)} characters")
                    return text, metadata
                else:
                    logger.warning(f"OCR quality too low with {engine_name}")
                    
            except Exception as e:
                logger.error(f"OCR failed with {engine_name}: {str(e)}")
                continue
        
        raise OCRError("All OCR engines failed to extract text")
    
    async def _extract_pdf_pages(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract pages from PDF as images for OCR."""
        pages_data = []
        
        try:
            doc = fitz.open(str(file_path))
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Convert page to image
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better OCR
                pix = page.get_pixmap(matrix=mat)
                img_data = pix.tobytes("png")
                image = Image.open(io.BytesIO(img_data))
                
                pages_data.append({
                    'image': image,
                    'page_num': page_num + 1,
                    'width': pix.width,
                    'height': pix.height
                })
            
            doc.close()
            
        except Exception as e:
            raise OCRError(f"Failed to extract PDF pages: {str(e)}") from e
        
        return pages_data
    
    async def _extract_docx_text(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract text directly from DOCX without OCR."""
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_parts.append(cell.text)
            
            full_text = '\n'.join(text_parts)
            
            metadata = {
                'extraction_method': 'direct',
                'total_pages': len(doc.sections),
                'confidence': 1.0,  # Direct extraction is 100% confident
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return full_text, metadata
            
        except Exception as e:
            raise OCRError(f"Failed to extract DOCX text: {str(e)}") from e
    
    async def _run_ocr_engine(self, engine_name: str, pages_data: List[Dict]) -> Tuple[str, Dict]:
        """Run specific OCR engine on pages."""
        engine_func = self.engines[engine_name]
        
        all_text = []
        all_confidences = []
        
        for page_data in pages_data:
            page_text, confidence = await engine_func(page_data['image'])
            all_text.append(page_text)
            all_confidences.append(confidence)
        
        full_text = '\n\n'.join(all_text)
        avg_confidence = sum(all_confidences) / len(all_confidences) if all_confidences else 0.0
        
        metadata = {
            'confidence': avg_confidence,
            'page_confidences': all_confidences,
            'engine': engine_name
        }
        
        return full_text, metadata
    
    @retry_on_failure(max_retries=2)
    async def _tesseract_ocr(self, image: Image.Image) -> Tuple[str, float]:
        """Tesseract OCR implementation."""
        try:
            # Configure Tesseract for better accuracy
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,;:!?()[]{}"\'-/$%&*+='
            
            # Extract text
            text = pytesseract.image_to_string(image, config=custom_config)
            
            # Get confidence data
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return text.strip(), avg_confidence / 100.0  # Convert to 0-1 scale
            
        except Exception as e:
            raise OCRError(f"Tesseract OCR failed: {str(e)}") from e
    
    async def _azure_ocr(self, image: Image.Image) -> Tuple[str, float]:
        """Azure Cognitive Services OCR implementation."""
        if not (self.config.azure_endpoint and self.config.azure_key):
            raise OCRError("Azure OCR credentials not configured")
        
        try:
            from azure.cognitiveservices.vision.computervision import ComputerVisionClient
            from azure.cognitiveservices.vision.computervision.models import OperationStatusCodes
            from msrest.authentication import CognitiveServicesCredentials
            
            # Initialize client
            credentials = CognitiveServicesCredentials(self.config.azure_key)
            client = ComputerVisionClient(self.config.azure_endpoint, credentials)
            
            # Convert image to bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            # Submit OCR request
            read_response = client.read_in_stream(img_byte_arr, raw=True)
            operation_id = read_response.headers["Operation-Location"].split("/")[-1]
            
            # Wait for completion
            while True:
                read_result = client.get_read_result(operation_id)
                if read_result.status not in ['notStarted', 'running']:
                    break
                await asyncio.sleep(1)
            
            # Extract text and confidence
            text_parts = []
            confidences = []
            
            if read_result.status == OperationStatusCodes.succeeded:
                for text_result in read_result.analyze_result.read_results:
                    for line in text_result.lines:
                        text_parts.append(line.text)
                        confidences.append(line.appearance.style.confidence if line.appearance else 0.9)
            
            full_text = '\n'.join(text_parts)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return full_text, avg_confidence
            
        except Exception as e:
            raise OCRError(f"Azure OCR failed: {str(e)}") from e
    
    async def _aws_ocr(self, image: Image.Image) -> Tuple[str, float]:
        """AWS Textract OCR implementation."""
        try:
            import boto3
            
            # Initialize Textract client
            textract = boto3.client('textract', region_name=self.config.aws_region)
            
            # Convert image to bytes
            img_byte_arr = io.BytesIO()
            image.save(img_byte_arr, format='PNG')
            img_bytes = img_byte_arr.getvalue()
            
            # Call Textract
            response = textract.detect_document_text(
                Document={'Bytes': img_bytes}
            )
            
            # Extract text and confidence
            text_parts = []
            confidences = []
            
            for item in response['Blocks']:
                if item['BlockType'] == 'LINE':
                    text_parts.append(item['Text'])
                    confidences.append(item['Confidence'] / 100.0)
            
            full_text = '\n'.join(text_parts)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return full_text, avg_confidence
            
        except Exception as e:
            raise OCRError(f"AWS Textract failed: {str(e)}") from e
    
    async def _easyocr_ocr(self, image: Image.Image) -> Tuple[str, float]:
        """EasyOCR implementation as fallback."""
        try:
            import easyocr
            
            # Initialize EasyOCR reader
            reader = easyocr.Reader(['en'])
            
            # Convert PIL Image to numpy array
            import numpy as np
            img_array = np.array(image)
            
            # Perform OCR
            results = reader.readtext(img_array)
            
            # Extract text and confidence
            text_parts = []
            confidences = []
            
            for (bbox, text, confidence) in results:
                text_parts.append(text)
                confidences.append(confidence)
            
            full_text = '\n'.join(text_parts)
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
            
            return full_text, avg_confidence
            
        except Exception as e:
            raise OCRError(f"EasyOCR failed: {str(e)}") from e
    
    def _get_document_type(self, file_path: Path) -> str:
        """Determine document type from file extension."""
        extension = file_path.suffix.lower()
        
        if extension == '.pdf':
            return 'pdf'
        elif extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.tif']:
            return 'image'
        elif extension in ['.docx', '.doc']:
            return 'docx'
        elif extension == '.txt':
            return 'text'
        else:
            raise OCRError(f"Unsupported file type: {extension}")
    
    def _validate_extraction(self, text: str, metadata: Dict) -> bool:
        """Validate OCR extraction quality."""
        # Check minimum text length
        if len(text.strip()) < 10:
            return False
        
        # Check confidence threshold
        confidence = metadata.get('confidence', 0.0)
        if confidence < self.config.confidence_threshold:
            return False
        
        # Check for reasonable character distribution
        alphanumeric = sum(c.isalnum() for c in text)
        if len(text) > 0 and alphanumeric / len(text) < 0.5:
            return False
        
        return True
    
    def _check_available_engines(self) -> Dict[str, bool]:
        """Check which OCR engines are available."""
        available = {}
        
        # Check Tesseract
        try:
            pytesseract.get_tesseract_version()
            available['tesseract'] = True
        except:
            available['tesseract'] = False
        
        # Check Azure
        try:
            from azure.cognitiveservices.vision.computervision import ComputerVisionClient
            available['azure'] = bool(self.config.azure_endpoint and self.config.azure_key)
        except ImportError:
            available['azure'] = False
        
        # Check AWS
        try:
            import boto3
            available['aws'] = True
        except ImportError:
            available['aws'] = False
        
        # Check EasyOCR
        try:
            import easyocr
            available['easyocr'] = True
        except ImportError:
            available['easyocr'] = False
        
        return available
