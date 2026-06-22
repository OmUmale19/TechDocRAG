"""
Document layout analysis for structure understanding.
Extracts layout elements like headers, tables, images, and their relationships.
"""

import cv2
import numpy as np
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from PIL import Image
import fitz  # PyMuPDF

from ..core.types import BoundingBox
from ..core.config import ProcessingConfig
from ..utils.logging import get_logger, performance_log
from ..utils.exceptions import ProcessingError

logger = get_logger(__name__)


class LayoutAnalyzer:
    """
    Advanced layout analysis for document structure understanding.
    
    Key Capabilities:
    1. Table detection and extraction
    2. Header/footer identification
    3. Multi-column layout handling
    4. Reading order determination
    5. Image/figure detection
    """
    
    def __init__(self, config: ProcessingConfig):
        self.config = config
        self.min_table_confidence = 0.7
        self.min_text_height = 10
        
    @performance_log("layout_analysis")
    async def analyze(self, file_path: Path, ocr_metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """
        Analyze document layout and extract structural elements.
        
        Args:
            file_path: Path to document file
            ocr_metadata: OCR metadata with confidence scores
            
        Returns:
            List of layout elements with type, text, and position
        """
        logger.info(f"Analyzing layout: {file_path}")
        
        try:
            file_extension = file_path.suffix.lower()
            
            if file_extension == '.pdf':
                elements = await self._analyze_pdf_layout(file_path)
            elif file_extension in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff']:
                elements = await self._analyze_image_layout(file_path)
            elif file_extension in ['.docx', '.doc']:
                elements = await self._analyze_docx_layout(file_path)
            else:
                # Fallback to basic text structure
                elements = await self._analyze_text_layout(file_path)
            
            # Post-process and validate elements
            elements = self._post_process_elements(elements)
            
            logger.info(f"Layout analysis completed: {len(elements)} elements found")
            return elements
            
        except Exception as e:
            logger.error(f"Layout analysis failed: {str(e)}")
            # Return basic structure on failure
            return [{'type': 'paragraph', 'text': '', 'confidence': 0.5, 'page': 1}]
    
    async def _analyze_pdf_layout(self, file_path: Path) -> List[Dict[str, Any]]:
        """Analyze PDF layout using PyMuPDF."""
        elements = []
        
        try:
            doc = fitz.open(str(file_path))
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get text blocks with position information
                blocks = page.get_text("dict")
                
                for block in blocks["blocks"]:
                    if "lines" in block:  # Text block
                        element = self._process_text_block(block, page_num + 1)
                        if element:
                            elements.append(element)
                    else:  # Image block
                        element = self._process_image_block(block, page_num + 1)
                        if element:
                            elements.append(element)
                
                # Detect tables using geometric analysis
                table_elements = self._detect_pdf_tables(page, page_num + 1)
                elements.extend(table_elements)
            
            doc.close()
            
        except Exception as e:
            raise ProcessingError(f"PDF layout analysis failed: {str(e)}") from e
        
        return elements
    
    def _process_text_block(self, block: Dict, page_num: int) -> Optional[Dict[str, Any]]:
        """Process a text block from PDF."""
        bbox = block["bbox"]
        
        # Extract text from all lines in block
        text_parts = []
        font_sizes = []
        
        for line in block["lines"]:
            for span in line["spans"]:
                text_parts.append(span["text"])
                font_sizes.append(span["size"])
        
        text = " ".join(text_parts).strip()
        if not text:
            return None
        
        # Determine element type based on formatting
        avg_font_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
        element_type = self._classify_text_element(text, avg_font_size, bbox)
        
        return {
            'type': element_type,
            'text': text,
            'bbox': BoundingBox(bbox[0], bbox[1], bbox[2], bbox[3], page_num),
            'page': page_num,
            'confidence': 0.9,  # High confidence for PDF text
            'font_size': avg_font_size,
            'metadata': {
                'font_sizes': font_sizes,
                'line_count': len(block["lines"])
            }
        }
    
    def _process_image_block(self, block: Dict, page_num: int) -> Dict[str, Any]:
        """Process an image block from PDF."""
        bbox = block["bbox"]
        
        return {
            'type': 'image',
            'text': '',
            'bbox': BoundingBox(bbox[0], bbox[1], bbox[2], bbox[3], page_num),
            'page': page_num,
            'confidence': 0.95,
            'metadata': {
                'width': bbox[2] - bbox[0],
                'height': bbox[3] - bbox[1]
            }
        }
    
    def _detect_pdf_tables(self, page, page_num: int) -> List[Dict[str, Any]]:
        """Detect tables in PDF page using geometric analysis."""
        tables = []
        
        try:
            # Get all horizontal and vertical lines
            drawings = page.get_drawings()
            h_lines = []
            v_lines = []
            
            for drawing in drawings:
                for item in drawing["items"]:
                    if item[0] == "l":  # Line
                        x1, y1, x2, y2 = item[1:5]
                        if abs(y1 - y2) < 2:  # Horizontal line
                            h_lines.append((x1, y1, x2, y2))
                        elif abs(x1 - x2) < 2:  # Vertical line
                            v_lines.append((x1, y1, x2, y2))
            
            # Find table regions based on line intersections
            if len(h_lines) >= 2 and len(v_lines) >= 2:
                table_regions = self._find_table_regions(h_lines, v_lines)
                
                for region in table_regions:
                    # Extract text within table region
                    table_text = self._extract_table_text(page, region)
                    
                    if table_text:
                        tables.append({
                            'type': 'table',
                            'text': table_text,
                            'bbox': BoundingBox(region[0], region[1], region[2], region[3], page_num),
                            'page': page_num,
                            'confidence': self.min_table_confidence,
                            'metadata': {
                                'rows': table_text.count('\n') + 1,
                                'estimated_columns': table_text.count('\t') + 1
                            }
                        })
        
        except Exception as e:
            logger.warning(f"Table detection failed: {str(e)}")
        
        return tables
    
    def _find_table_regions(self, h_lines: List[Tuple], v_lines: List[Tuple]) -> List[Tuple]:
        """Find rectangular table regions from line intersections."""
        regions = []
        
        # Sort lines
        h_lines.sort(key=lambda x: x[1])  # Sort by y-coordinate
        v_lines.sort(key=lambda x: x[0])  # Sort by x-coordinate
        
        # Find rectangular regions formed by line intersections
        for i in range(len(h_lines) - 1):
            for j in range(len(v_lines) - 1):
                h1 = h_lines[i]
                h2 = h_lines[i + 1]
                v1 = v_lines[j]
                v2 = v_lines[j + 1]
                
                # Check if lines form a rectangle
                x1, y1 = max(h1[0], v1[0]), h1[1]
                x2, y2 = min(h1[2], v2[0]), h2[1]
                
                if x2 > x1 and y2 > y1:
                    regions.append((x1, y1, x2, y2))
        
        return regions
    
    def _extract_table_text(self, page, region: Tuple) -> str:
        """Extract text from a table region."""
        rect = fitz.Rect(region)
        text = page.get_text("text", clip=rect)
        
        # Convert to tab-separated format for better structure
        lines = text.strip().split('\n')
        processed_lines = []
        
        for line in lines:
            if line.strip():
                # Simple heuristic to detect columns (multiple spaces)
                cells = re.split(r'\s{2,}', line.strip())
                processed_lines.append('\t'.join(cells))
        
        return '\n'.join(processed_lines)
    
    async def _analyze_image_layout(self, file_path: Path) -> List[Dict[str, Any]]:
        """Analyze layout of image documents using computer vision."""
        elements = []
        
        try:
            # Load image
            image = cv2.imread(str(file_path))
            if image is None:
                raise ProcessingError(f"Could not load image: {file_path}")
            
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Detect text regions using MSER (Maximally Stable Extremal Regions)
            mser = cv2.MSER_create()
            regions, _ = mser.detectRegions(gray)
            
            # Group regions into text blocks
            text_blocks = self._group_text_regions(regions, image.shape)
            
            for i, block in enumerate(text_blocks):
                elements.append({
                    'type': 'text_block',
                    'text': '',  # Text will be filled by OCR
                    'bbox': BoundingBox(block[0], block[1], block[2], block[3], 1),
                    'page': 1,
                    'confidence': 0.7,
                    'metadata': {'block_index': i}
                })
            
            # Detect horizontal lines (potential table borders)
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (40, 1))
            horizontal_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, horizontal_kernel)
            
            # Detect vertical lines
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 40))
            vertical_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, vertical_kernel)
            
            # Find table regions from line intersections
            table_regions = self._find_image_tables(horizontal_lines, vertical_lines)
            
            for region in table_regions:
                elements.append({
                    'type': 'table',
                    'text': '',
                    'bbox': BoundingBox(region[0], region[1], region[2], region[3], 1),
                    'page': 1,
                    'confidence': 0.6,
                    'metadata': {'detected_from_lines': True}
                })
        
        except Exception as e:
            logger.error(f"Image layout analysis failed: {str(e)}")
            # Return basic single-block structure
            elements = [{
                'type': 'paragraph',
                'text': '',
                'bbox': None,
                'page': 1,
                'confidence': 0.5
            }]
        
        return elements
    
    def _group_text_regions(self, regions: List, image_shape: Tuple) -> List[Tuple]:
        """Group MSER regions into coherent text blocks."""
        if not regions:
            return []
        
        # Convert regions to bounding boxes
        boxes = []
        for region in regions:
            x_coords = region[:, 0]
            y_coords = region[:, 1]
            x1, y1 = np.min(x_coords), np.min(y_coords)
            x2, y2 = np.max(x_coords), np.max(y_coords)
            
            # Filter out very small or very large regions
            width, height = x2 - x1, y2 - y1
            if 10 < width < image_shape[1] * 0.8 and 5 < height < image_shape[0] * 0.8:
                boxes.append((x1, y1, x2, y2))
        
        # Group nearby boxes
        grouped_boxes = self._merge_nearby_boxes(boxes)
        
        return grouped_boxes
    
    def _merge_nearby_boxes(self, boxes: List[Tuple], distance_threshold: int = 20) -> List[Tuple]:
        """Merge nearby bounding boxes into larger text blocks."""
        if not boxes:
            return []
        
        merged = []
        used = set()
        
        for i, box1 in enumerate(boxes):
            if i in used:
                continue
            
            group = [box1]
            used.add(i)
            
            # Find nearby boxes
            for j, box2 in enumerate(boxes):
                if j in used:
                    continue
                
                if self._boxes_are_close(box1, box2, distance_threshold):
                    group.append(box2)
                    used.add(j)
            
            # Merge all boxes in group
            if group:
                x1 = min(box[0] for box in group)
                y1 = min(box[1] for box in group)
                x2 = max(box[2] for box in group)
                y2 = max(box[3] for box in group)
                merged.append((x1, y1, x2, y2))
        
        return merged
    
    def _boxes_are_close(self, box1: Tuple, box2: Tuple, threshold: int) -> bool:
        """Check if two bounding boxes are close enough to merge."""
        x1_1, y1_1, x2_1, y2_1 = box1
        x1_2, y1_2, x2_2, y2_2 = box2
        
        # Calculate distance between box centers
        center1_x, center1_y = (x1_1 + x2_1) / 2, (y1_1 + y2_1) / 2
        center2_x, center2_y = (x1_2 + x2_2) / 2, (y1_2 + y2_2) / 2
        
        distance = np.sqrt((center1_x - center2_x)**2 + (center1_y - center2_y)**2)
        
        return distance < threshold
    
    def _find_image_tables(self, h_lines_img, v_lines_img) -> List[Tuple]:
        """Find table regions from detected lines in image."""
        tables = []
        
        try:
            # Find contours in line images
            h_contours, _ = cv2.findContours(h_lines_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            v_contours, _ = cv2.findContours(v_lines_img, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Extract line coordinates
            h_lines = []
            v_lines = []
            
            for contour in h_contours:
                x, y, w, h = cv2.boundingRect(contour)
                if w > 50:  # Minimum line length
                    h_lines.append((x, y, x + w, y))
            
            for contour in v_contours:
                x, y, w, h = cv2.boundingRect(contour)
                if h > 50:  # Minimum line length
                    v_lines.append((x, y, x, y + h))
            
            # Find intersections to identify table regions
            if len(h_lines) >= 2 and len(v_lines) >= 2:
                tables = self._find_table_regions(h_lines, v_lines)
        
        except Exception as e:
            logger.warning(f"Image table detection failed: {str(e)}")
        
        return tables
    
    async def _analyze_docx_layout(self, file_path: Path) -> List[Dict[str, Any]]:
        """Analyze DOCX document structure."""
        elements = []
        
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            element_index = 0
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    element_type = self._classify_docx_paragraph(para)
                    
                    elements.append({
                        'type': element_type,
                        'text': para.text,
                        'page': 1,  # DOCX doesn't have page info easily accessible
                        'confidence': 0.95,
                        'metadata': {
                            'style': para.style.name if para.style else 'Normal',
                            'element_index': element_index
                        }
                    })
                    element_index += 1
            
            # Process tables
            for table in doc.tables:
                table_text = self._extract_docx_table_text(table)
                
                elements.append({
                    'type': 'table',
                    'text': table_text,
                    'page': 1,
                    'confidence': 0.9,
                    'metadata': {
                        'rows': len(table.rows),
                        'columns': len(table.columns) if table.rows else 0,
                        'element_index': element_index
                    }
                })
                element_index += 1
        
        except Exception as e:
            logger.error(f"DOCX layout analysis failed: {str(e)}")
            return [{'type': 'paragraph', 'text': '', 'page': 1, 'confidence': 0.5}]
        
        return elements
    
    def _classify_docx_paragraph(self, paragraph) -> str:
        """Classify DOCX paragraph type based on style and content."""
        style_name = paragraph.style.name.lower() if paragraph.style else ''
        text = paragraph.text.strip()
        
        if 'heading' in style_name or 'title' in style_name:
            return 'header'
        elif 'list' in style_name or text.startswith(('-', '*', '•')):
            return 'list'
        elif len(text) < 100 and text.isupper():
            return 'header'
        else:
            return 'paragraph'
    
    def _extract_docx_table_text(self, table) -> str:
        """Extract text from DOCX table in structured format."""
        rows = []
        
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append('\t'.join(cells))
        
        return '\n'.join(rows)
    
    async def _analyze_text_layout(self, file_path: Path) -> List[Dict[str, Any]]:
        """Basic text layout analysis for unsupported formats."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Simple paragraph detection
            paragraphs = content.split('\n\n')
            elements = []
            
            for i, para in enumerate(paragraphs):
                if para.strip():
                    elements.append({
                        'type': 'paragraph',
                        'text': para.strip(),
                        'page': 1,
                        'confidence': 0.8,
                        'metadata': {'paragraph_index': i}
                    })
            
            return elements
            
        except Exception as e:
            logger.error(f"Text layout analysis failed: {str(e)}")
            return [{'type': 'paragraph', 'text': '', 'page': 1, 'confidence': 0.5}]
    
    def _classify_text_element(self, text: str, font_size: float, bbox: Tuple) -> str:
        """Classify text element type based on content and formatting."""
        text_lower = text.lower().strip()
        
        # Header detection
        if font_size > 14 or (len(text) < 100 and text.isupper()):
            return 'header'
        
        # List detection
        if re.match(r'^\s*[•\-\*\d+\.]\s+', text):
            return 'list'
        
        # Table detection (multiple columns of data)
        if '\t' in text or re.search(r'\d+\.\d+\s+\d+\.\d+', text):
            return 'table'
        
        # Footer detection (small text at bottom)
        if font_size < 10:
            return 'footer'
        
        return 'paragraph'
    
    def _post_process_elements(self, elements: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Post-process layout elements for consistency and quality."""
        processed = []
        
        for element in elements:
            # Skip empty elements
            if not element.get('text', '').strip() and element.get('type') != 'image':
                continue
            
            # Ensure required fields
            element.setdefault('confidence', 0.5)
            element.setdefault('page', 1)
            element.setdefault('metadata', {})
            
            # Validate bounding box
            if element.get('bbox') and not isinstance(element['bbox'], BoundingBox):
                bbox_data = element['bbox']
                if isinstance(bbox_data, (list, tuple)) and len(bbox_data) >= 4:
                    element['bbox'] = BoundingBox(
                        bbox_data[0], bbox_data[1], bbox_data[2], bbox_data[3],
                        element['page']
                    )
            
            processed.append(element)
        
        # Sort elements by reading order (top to bottom, left to right)
        processed.sort(key=lambda x: (
            x['page'],
            x['bbox'].y1 if x.get('bbox') else 0,
            x['bbox'].x1 if x.get('bbox') else 0
        ))
        
        return processed


# Import missing module
import re
